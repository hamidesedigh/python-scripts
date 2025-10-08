# -*- coding: utf-8 -*-
"""
@author: Hamideh
@created: Tues Oct 7 16:00:00 2025

Program Description:
-------------------------
Find and Replace Numbers in a Word Document
--------------------------------------------------------------
Covers:
✔ python-docx library
✔ Reading paragraphs
✔ Replacing specific patterns (e.g., numbers)
✔ Saving updated file
"""

from docx import Document
import re

input_file = "sample.docx"
output_file = "updated.docx"

# Create a sample DOCX
doc = Document()
doc.add_paragraph("Invoice Number: 12345")
doc.add_paragraph("Total Amount: 56789")
doc.save(input_file)
print("✅ Sample DOCX created.")

# --- Read and Replace Numbers ---
doc = Document(input_file)
for para in doc.paragraphs:
    para.text = re.sub(r"\d+", "[NUMBER]", para.text)

doc.save(output_file)
print("✅ Numbers replaced and new DOCX saved.")

"""
Quick Summary:
--------------
✔ python-docx → Read and write Word documents
✔ re.sub() → Replace text patterns (e.g., digits)
"""
