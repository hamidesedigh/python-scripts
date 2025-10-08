# -*- coding: utf-8 -*-
"""
@author: Hamideh
@created: Wed Jun 25 06:51:20 2025

Program Description:
-------------------------
Demonstration of Word (.docx) File Handling using python-docx
--------------------------------------------------------------
Covers:
✔ Reading and writing Word files
✔ Modifying text content programmatically
✔ Detecting and processing digit sequences
✔ Extracting and summing English/Persian numbers

Requirements:
-------------
pip install python-docx
"""

from docx import Document

# ============================================================
# --- Mini Project 1: Modify Digit Sequences in DOCX ---
# ============================================================

def process_text(text: str) -> tuple[str, int]:
    """
    Modify digit sequences by replacing consecutive digits with the last one.

    Parameters:
        text (str): The paragraph text.

    Returns:
        tuple[str, int]: (Processed text, number of fixes applied)
    """
    new_text = []
    numbers_done = 0
    fixes = 0

    for char in text:
        if char.isdigit():
            fixes += 1
            if numbers_done == 0:
                new_text.append(char)
            else:
                # Replace the previous digit in the same sequence
                new_text[-numbers_done] = char
            numbers_done += 1
        else:
            new_text.append(char)
            numbers_done = 0

    return "".join(new_text), fixes


def modify_docx_digits(input_file: str, output_file: str):
    """
    Open a DOCX file, modify its digits, and save a revised version.
    """
    print("\n--- Running Mini Project 1: Modify Digit Sequences ---")
    doc = Document(input_file)
    total_fixes = 0

    for para in doc.paragraphs:
        new_text, fixes = process_text(para.text)
        para.text = new_text
        total_fixes += fixes

    doc.save(output_file)
    print(f"✅ Processing complete. Total fixes applied: {total_fixes}")
    print(f"💾 Revised file saved as: {output_file}")


# ============================================================
# --- Mini Project 2: Extract and Sum Numbers (Eng/Persian) ---
# ============================================================

# Persian-to-English digit map
PERSIAN_DIGITS = str.maketrans("۰۱۲۳۴۵۶۷۸۹", "0123456789")

def extract_and_sum_numbers(text: str) -> tuple[float, int]:
    """
    Extract numbers (English or Persian) from text and calculate their total.

    Parameters:
        text (str): Input paragraph text.

    Returns:
        tuple[float, int]: (Sum of numbers, count of numbers)
    """
    num_str = ""
    total = 0.0
    count = 0

    for char in text:
        if char.isdigit() or char in "۰۱۲۳۴۵۶۷۸۹":
            num_str += char.translate(PERSIAN_DIGITS)
        elif char in ".٫" and num_str:
            num_str += "."
        elif num_str and num_str != ".":
            total += float(num_str)
            count += 1
            num_str = ""

    # Handle trailing number
    if num_str and num_str != ".":
        total += float(num_str)
        count += 1

    return total, count


def sum_docx_numbers(input_file: str):
    """
    Open a DOCX file, extract all numeric values, and compute their sum.
    """
    print("\n--- Running Mini Project 2: Extract and Sum Numbers ---")
    doc = Document(input_file)

    grand_total = 0.0
    grand_count = 0

    for para in doc.paragraphs:
        subtotal, count = extract_and_sum_numbers(para.text)
        grand_total += subtotal
        grand_count += count

    print(f"✅ Total sum of numbers: {grand_total}")
    print(f"🔢 Total numbers found: {grand_count}")


# ============================================================
# --- MAIN EXECUTION ---
# ============================================================

if __name__ == "__main__":
    input_doc = "FinalProject.docx"
    output_doc = "Revised.docx"

    # Run both mini projects
    modify_docx_digits(input_doc, output_doc)
    sum_docx_numbers(input_doc)

# ============================================================
# --- Quick Summary ---
# ============================================================
"""
Quick Summary:
--------------
✔ Document(input_file) → Open a .docx file.
✔ doc.paragraphs → Access list of paragraphs.
✔ para.text → Read or modify text content.
✔ doc.save(output_file) → Save document changes.

Mini Projects Recap:
1️⃣ Modify Digits → Iterates through paragraphs and adjusts digit sequences.
2️⃣ Sum Numbers → Finds English/Persian numbers and calculates their total.

Tips:
- Use regex for more complex text extraction.
- python-docx can also handle tables, runs, and styles.
"""
